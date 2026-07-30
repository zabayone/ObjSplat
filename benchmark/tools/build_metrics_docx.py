#!/usr/bin/env python3
"""Build the ObjSplat benchmark metric catalogue DOCX."""
from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from benchmark import schemas

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6B78"
HEADER_FILL = "E8EEF5"
ALT_FILL = "F7F9FC"
BORDER = "B7C4D2"
CALLOUT = "F4F6F9"
CONTENT_WIDTH_DXA = 9360


SECTIONS = [
    ("Tempi di esecuzione", "stage_timings.csv", schemas.STAGE_COLUMNS),
    ("Campionamento delle risorse", "resource_samples.csv", schemas.RESOURCE_COLUMNS),
    ("Metriche per layer", "layer_metrics.csv", schemas.LAYER_COLUMNS),
    ("Segmentazione e copertura", "segmentation_metrics.csv", schemas.SEGMENTATION_COLUMNS),
    ("Qualità di ricostruzione", "reconstruction_metrics.csv", schemas.RECONSTRUCTION_COLUMNS),
    ("Prestazioni di rendering", "rendering_metrics.csv", schemas.RENDERING_COLUMNS),
    ("Località dell’editing", "editing_metrics.csv", schemas.EDITING_COLUMNS),
    ("Conservazione topologica e mood", "mood_metrics.csv", schemas.MOOD_COLUMNS),
]


DESCRIPTIONS = {
    "experiment": "Nome stabile dell’esperimento risolto dalla configurazione.",
    "scene": "Identificatore della scena/ERP.",
    "run_id": "Identificatore univoco del singolo run; impedisce la commistione fra ripetizioni.",
    "variant": "Variante valutata: layered, refined, night, monolithic o singolo layer.",
    "stage": "Nome strutturato dello stadio della pipeline.",
    "parent_stage": "Stadio contenitore, quando la misura è annidata.",
    "started_at": "Timestamp UTC ISO-8601 di inizio.",
    "ended_at": "Timestamp UTC ISO-8601 di fine.",
    "wall_seconds": "Durata reale trascorsa misurata con clock monotono.",
    "cpu_seconds": "Tempo CPU consumato dal processo durante lo stadio.",
    "status": "Esito: success, failed, skipped, partial_success o unavailable.",
    "rss_before_bytes": "Resident Set Size del processo prima dello stadio.",
    "rss_after_bytes": "Resident Set Size del processo dopo lo stadio.",
    "system_available_before_bytes": "Memoria di sistema disponibile prima dello stadio.",
    "system_available_after_bytes": "Memoria di sistema disponibile dopo lo stadio.",
    "peak_sampled_rss_bytes": "Massimo RSS osservato dal campionatore durante il run/stadio.",
    "input_points": "Numero di punti/Gaussiane forniti in ingresso allo stadio.",
    "output_gaussians": "Numero di Gaussiane prodotte dallo stadio.",
    "iterations": "Iterazioni configurate o effettivamente completate.",
    "frames": "Numero di frame coinvolti nello stadio.",
    "layer_index": "Indice numerico del layer associato.",
    "exception_type": "Classe dell’eccezione in caso di fallimento.",
    "exception_message": "Messaggio diagnostico dell’eccezione.",
    "seconds_per_iteration": "Durata normalizzata per iterazione.",
    "seconds_per_frame": "Durata normalizzata per frame.",
    "seconds_per_million_input_points": "Durata normalizzata per milione di punti in ingresso.",
    "seconds_per_million_output_gaussians": "Durata normalizzata per milione di Gaussiane finali.",
    "timestamp": "Timestamp UTC del campione periodico.",
    "elapsed_seconds": "Tempo trascorso dall’avvio del campionatore/run.",
    "process_rss_bytes": "RSS aggregato del benchmark e dei processi figli.",
    "process_vms_bytes": "Memoria virtuale aggregata del processo e dei figli.",
    "system_total_bytes": "Memoria fisica/unificata totale rilevata dal sistema.",
    "system_available_bytes": "Memoria che il sistema stima immediatamente disponibile.",
    "system_used_bytes": "Memoria di sistema attualmente utilizzata.",
    "system_used_percent": "Percentuale di memoria di sistema utilizzata.",
    "swap_total_bytes": "Capacità swap totale.",
    "swap_used_bytes": "Swap utilizzata al momento del campione.",
    "process_cpu_percent": "Utilizzo CPU istantaneo del processo campionato.",
    "semantic_label": "Etichetta semantica aggregata del layer.",
    "instance_ids": "Lista JSON degli ID di istanza contenuti nel layer.",
    "confidence_count": "Numero di confidence GroundingDINO disponibili.",
    "confidence_mean": "Confidence GroundingDINO media delle istanze del layer.",
    "confidence_min": "Confidence GroundingDINO minima.",
    "confidence_max": "Confidence GroundingDINO massima.",
    "mask_area_pixels": "Area della maschera ERP del layer.",
    "mask_coverage_percent": "Frazione percentuale dell’ERP assegnata al layer.",
    "connected_components": "Numero di componenti connesse nella maschera ERP.",
    "projected_3d_points": "Punti 3D assegnati al layer dopo proiezione e filtraggio.",
    "training_frames": "Frame prospettici con supervisione non vuota per il layer.",
    "total_supervised_pixels": "Somma dei pixel supervisionati su tutti i frame del layer.",
    "mean_supervised_pixels_per_frame": "Media dei pixel supervisionati per frame.",
    "training_iterations": "Budget di training attribuito al layer.",
    "training_time_seconds": "Tempo strutturato di training del layer.",
    "initial_gaussians": "Numero iniziale di punti/Gaussiane nel PLY di training.",
    "final_gaussians": "Numero di Gaussiane nel PLY addestrato.",
    "ply_size_bytes": "Dimensione su disco del PLY del layer.",
    "percent_final_scene_gaussians": "Quota del numero finale di Gaussiane rappresentata dal layer.",
    "reason": "Motivo di indisponibilità, skip o fallimento.",
    "target": "Etichetta, layer, variante o file oggetto della misura.",
    "metric_scope": "Ambito: intrinsic oppure ground_truth.",
    "iou": "Intersection over Union fra maschera predetta e annotazione.",
    "dice": "Coefficiente Dice fra maschera predetta e annotazione.",
    "precision": "Precision pixel-wise: TP/(TP+FP).",
    "recall": "Recall pixel-wise: TP/(TP+FN).",
    "boundary_fscore": "F-score dei bordi con tolleranza morfologica di due pixel.",
    "false_positive_pixels": "Pixel predetti come target ma assenti nel ground truth.",
    "false_negative_pixels": "Pixel target non recuperati dalla predizione.",
    "thin_structure_iou": "IoU opzionale per strutture sottili, solo con annotazione dedicata.",
    "coverage_percent": "Copertura intrinseca della maschera sul panorama.",
    "background_percent": "Percentuale dell’ERP assegnata al background.",
    "unassigned_percent": "Percentuale ERP non assegnata ad alcun layer.",
    "overlap_before_pixels": "Pixel coperti da più maschere prima della risoluzione dei conflitti, se persistiti.",
    "overlap_after_pixels": "Pixel ancora sovrapposti dopo la risoluzione dei conflitti.",
    "seam_crossing": "Indica se la maschera attraversa il seam orizzontale dell’ERP.",
    "source": "File sorgente della maschera o annotazione.",
    "view_id": "Indice deterministico della vista prospettica.",
    "split": "Appartenenza allo split: held_out o stato di indisponibilità.",
    "theta_deg": "Orientamento orizzontale della camera.",
    "phi_deg": "Orientamento verticale della camera.",
    "width": "Larghezza del render o riferimento.",
    "height": "Altezza del render o riferimento.",
    "psnr_db": "Peak Signal-to-Noise Ratio RGB; valori maggiori indicano errore quadratico minore.",
    "ssim": "Structural Similarity Index, se scikit-image è disponibile.",
    "lpips": "Distanza percettiva LPIPS; valori minori sono migliori.",
    "mae": "Errore assoluto medio RGB normalizzato in [0,1].",
    "foreground_psnr_db": "PSNR limitato alla regione foreground proiettata.",
    "foreground_ssim": "SSIM limitato alla regione foreground.",
    "background_psnr_db": "PSNR limitato alla regione background.",
    "background_ssim": "SSIM limitato alla regione background.",
    "render_seconds": "Durata del singolo render di valutazione.",
    "note": "Nota metodologica o limitazione associata alla riga.",
    "warmup_frames": "Frame di warm-up esclusi dalle statistiche steady-state.",
    "measured_frames": "Frame cronometrati dopo il warm-up.",
    "cold_start_seconds": "Prima renderizzazione, inclusi caricamento/compilazione iniziale.",
    "mean_ms": "Latenza media steady-state per frame.",
    "median_ms": "Latenza mediana per frame.",
    "p90_ms": "90º percentile della latenza per frame.",
    "p95_ms": "95º percentile della latenza per frame.",
    "average_fps": "FPS derivati dalla latenza media.",
    "minimum_fps": "FPS corrispondenti alla latenza massima osservata.",
    "gaussian_count": "Numero di Gaussiane della rappresentazione renderizzata.",
    "megapixels_per_second": "Throughput normalizzato per risoluzione.",
    "target_type": "Granularità dell’edit: layer oppure instance.",
    "target_id": "Indice layer o ID istanza rimosso.",
    "inside_changed_percent": "Pixel modificati all’interno della maschera target.",
    "outside_changed_percent": "Pixel modificati all’esterno della maschera target.",
    "outside_mae": "Errore assoluto medio fuori dalla regione target.",
    "outside_lpips": "Distanza percettiva fuori maschera, quando tecnicamente disponibile.",
    "edit_leakage_ratio": "Rapporto fra cambiamento medio esterno e interno; minore è migliore.",
    "edit_locality_score": "Quota del cambiamento concentrata nel target; maggiore è migliore.",
    "removed_gaussians": "Numero assoluto di Gaussiane rimosse.",
    "removed_gaussians_percent": "Percentuale di Gaussiane rimosse dalla scena.",
    "creation_seconds": "Tempo necessario per creare la rappresentazione editata.",
    "retraining_required": "Indica se l’edit ha richiesto nuovo training.",
    "edited_size_bytes": "Dimensione del PLY editato.",
    "day_variant": "PLY diurno usato come riferimento topologico.",
    "mood_variant": "PLY mood/notturno confrontato.",
    "correspondence_compatible": "Compatibilità di conteggio e schema prima del confronto elemento-per-elemento.",
    "day_gaussians": "Gaussiane presenti nella variante diurna.",
    "mood_gaussians": "Gaussiane presenti nella variante mood.",
    "gaussian_count_difference": "Differenza mood − day nel numero di Gaussiane.",
    "position_mean_abs": "Differenza assoluta media delle coordinate xyz.",
    "position_max_abs": "Massima differenza assoluta nelle coordinate xyz.",
    "scale_mean_abs": "Differenza assoluta media dei log-scale.",
    "scale_max_abs": "Massima differenza assoluta dei log-scale.",
    "rotation_mean_abs": "Differenza assoluta media delle componenti quaternion.",
    "rotation_max_abs": "Massima differenza assoluta delle componenti quaternion.",
    "opacity_mean_abs": "Differenza assoluta media dell’opacity logit.",
    "opacity_max_abs": "Massima differenza assoluta dell’opacity logit.",
    "label_difference_count": "Gaussiane il cui label intero è cambiato.",
    "sh_mean_abs": "Differenza assoluta media dei coefficienti SH.",
    "sh_max_abs": "Massima differenza assoluta dei coefficienti SH.",
    "nonappearance_changed_percent": "Gaussiane con proprietà non di apparenza cambiate oltre la tolleranza.",
    "analytic_fit_seconds": "Tempo dell’adattamento analitico ERP→SH.",
    "refinement_seconds": "Tempo dell’eventuale refinement SH-only.",
    "mood_ply_size_bytes": "Dimensione del PLY mood/notturno.",
    "target_erp_psnr_db": "PSNR rispetto all’ERP target mood, se è disponibile un protocollo valido.",
    "target_erp_ssim": "SSIM rispetto all’ERP target mood, se disponibile.",
    "circular_seam_mae": "Discontinuità cromatica media fra margine sinistro e destro dell’ERP.",
}


SYSTEM_ROWS = [
    ("collected_at", "timestamp", "Misurato", "Istante UTC di acquisizione dei metadati."),
    ("git.commit", "hash", "Misurato", "Commit Git del codice eseguito."),
    ("git.dirty", "boolean", "Misurato", "Indica modifiche non committate nel worktree."),
    ("git.status_porcelain", "testo", "Misurato", "Elenco compatto delle modifiche locali."),
    ("os.*", "testo", "Misurato", "Sistema operativo, release, versione e piattaforma."),
    ("python.*", "testo", "Misurato", "Versione ed eseguibile Python."),
    ("hardware.machine_model", "testo", "Misurato", "Modello macchina rilevato da sysctl/platform."),
    ("hardware.chip_model", "testo", "Misurato", "Modello Apple Silicon/CPU."),
    ("hardware.architecture", "testo", "Misurato", "Architettura del processore."),
    ("hardware.total_unified_memory_bytes", "byte", "Misurato", "Memoria fisica/unificata totale."),
    ("hardware.available_memory_before_bytes", "byte", "Misurato", "Memoria disponibile prima del run."),
    ("hardware.memory_semantics", "testo", "Metadato", "Distingue memoria unificata da CUDA VRAM."),
    ("packages.*", "versione", "Misurato", "Versioni delle dipendenze scientifiche rilevanti."),
    ("device.*", "testo/boolean", "Misurato", "Device richiesto e disponibilità MLX/Torch."),
    ("command_line", "lista", "Misurato", "Argomenti esatti usati per il run."),
    ("environment", "mappa", "Misurato", "Variabili d’ambiente rilevanti per pipeline/backend."),
    ("random_seeds", "intero", "Configurato", "Seed del benchmark e delle trasformazioni controllate."),
    ("panorama.*", "pixel/testo", "Misurato", "Path e dimensioni dell’ERP sorgente."),
    ("scene_name", "testo", "Configurato", "Nome della scena."),
    ("benchmark_configuration", "testo", "Configurato", "Nome della configurazione risolta."),
]


SUMMARY_ROWS = [
    ("status", "categoria", "Derivato", "Esito complessivo del run."),
    ("elapsed_seconds", "s", "Misurato", "Durata totale osservata dal recorder."),
    ("peak_process_rss_bytes", "byte", "Derivato", "Picco RSS dei campioni del run."),
    ("failed_stage", "testo", "Derivato", "Ultimo stadio fallito."),
    ("last_resource_sample", "record", "Misurato", "Ultimo campione disponibile prima della conclusione/failure."),
    ("completed_stages", "lista", "Misurato", "Stadi conclusi con il relativo esito."),
    ("outputs_successfully_generated", "lista", "Derivato", "File prodotti prima della fine o del fallimento."),
    ("usable_merged_ply", "boolean", "Derivato", "Presenza di una scena merged utilizzabile."),
    ("generated_training_views", "conteggio", "Misurato", "Viste prospettiche generate."),
    ("input_point_count", "conteggio", "Derivato", "Somma dei punti iniziali dei layer."),
    ("final_gaussian_count", "conteggio", "Derivato", "Somma delle Gaussiane finali dei layer."),
    ("disk_usage.*", "byte", "Derivato", "Spazio per scene, traindata, PLY, frame, mask, mood e totale."),
    ("layer_summary.*", "statistiche", "Derivato", "Conteggi, distribuzione, imbalance e correlazioni per layer."),
    ("segmentation_summary.*", "statistiche", "Derivato", "Copertura intrinseca, overlap, seam, discard e sky."),
    ("mood_variant_count", "conteggio", "Derivato", "Numero di varianti mood valide rilevate."),
]


def unit_for(name: str) -> str:
    if name.endswith("_bytes"):
        return "byte"
    if name.endswith("_seconds") or name.startswith("seconds_per"):
        return "s"
    if name.endswith("_ms"):
        return "ms"
    if name.endswith("_percent"):
        return "%"
    if name.endswith("_db"):
        return "dB"
    if name.endswith("_fps") or name == "average_fps":
        return "frame/s"
    if name in {"ssim", "iou", "dice", "precision", "recall", "boundary_fscore",
                "lpips", "mae", "outside_mae", "edit_leakage_ratio",
                "edit_locality_score"}:
        return "adimensionale"
    if "width" in name or "height" in name or name.endswith("_pixels"):
        return "pixel"
    if name in {"theta_deg", "phi_deg"}:
        return "gradi"
    if name.endswith("_count") or name.endswith("_gaussians") or name in {
        "iterations", "frames", "layer_index", "view_id", "training_frames",
        "connected_components", "projected_3d_points", "removed_gaussians",
    }:
        return "conteggio"
    if name in {"status", "reason", "note", "source", "variant", "target",
                "target_type", "target_id", "semantic_label", "split"}:
        return "testo"
    return "—"


def nature_for(name: str) -> str:
    if name in {"experiment", "scene", "run_id", "variant", "target", "target_type",
                "target_id", "semantic_label", "layer_index", "view_id", "split",
                "source", "day_variant", "mood_variant"}:
        return "Identificatore"
    if name in {"status", "reason", "note", "exception_type", "exception_message"}:
        return "Diagnostica"
    derived_tokens = (
        "percent", "mean", "median", "p90", "p95", "fps", "per_", "psnr",
        "ssim", "lpips", "mae", "iou", "dice", "precision", "recall",
        "fscore", "difference", "locality", "leakage", "coverage",
    )
    if any(token in name for token in derived_tokens):
        return "Derivato"
    return "Misurato"


def description_for(name: str) -> str:
    return DESCRIPTIONS.get(
        name,
        f"Campo `{name}` registrato nello schema stabile del benchmark.",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=11, color="000000", bold=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_end])
    set_run_font(run, size=9, color=MUTED)


def add_metric_table(doc, rows, widths=(2450, 1050, 1300, 4560)):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Campo", "Unità", "Natura", "Descrizione e interpretazione")
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9, color=INK, bold=True)
    set_repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, ALT_FILL)
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_run_font(run, size=8.5, color="000000", bold=(index == 0))
    set_table_geometry(table, list(widths))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_document(output: Path) -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("ObjSplat  |  Benchmark Metric Catalogue")
    set_run_font(run, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Editorial-cover header pattern, restrained for a technical reference.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("TECHNICAL REFERENCE")
    set_run_font(run, size=10, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("ObjSplat Benchmark Metrics")
    set_run_font(run, size=28, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run(
        "Catalogo delle misure, definizioni, unità e criteri di interpretazione"
    )
    set_run_font(run, size=13, color=DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(110)
    run = meta.add_run(f"Versione framework 1.0  •  {date.today().isoformat()}")
    set_run_font(run, size=10, color=MUTED, italic=True)
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lead.add_run(
        "Documento di riferimento per la metodologia sperimentale della tesi magistrale"
    )
    set_run_font(run, size=10.5, color=INK)

    doc.add_page_break()
    doc.add_heading("1. Scopo e criteri di lettura", level=1)
    doc.add_paragraph(
        "Il framework separa valori misurati direttamente, valori derivati, "
        "statistiche intrinseche e metriche che richiedono ground truth. I campi "
        "opzionali restano null quando la dipendenza o l’evidenza necessaria non "
        "è disponibile: non vengono imputati né presentati come accuratezza."
    )
    add_callout(
        doc,
        "Limite scientifico principale",
        "Le viste held-out sono proiezioni prospettiche della stessa ERP. Misurano "
        "fedeltà di ricostruzione panoramica, non accuratezza geometrica novel-view "
        "su camere indipendenti.",
    )
    for label, text in (
        ("Misurato", "Letto da clock, sistema, file, maschere, PLY o renderer."),
        ("Derivato", "Calcolato esclusivamente da valori misurati e con formula documentata."),
        ("Intrinseco", "Descrive copertura/consistenza senza ground truth; non è accuratezza."),
        ("Ground truth", "Disponibile solo con annotazioni manuali nel formato previsto."),
        ("Opzionale", "Può essere null per dipendenza assente o rappresentazione incompatibile."),
    ):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{label}: ")
        set_run_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(text)
        set_run_font(r)

    doc.add_heading("2. Metadati di sistema e riproducibilità", level=1)
    doc.add_paragraph(
        "Salvati in system_info.json prima dell’esperimento. Su Apple Silicon la "
        "memoria è unificata: RSS e memoria di sistema sono stime condivise CPU/GPU, "
        "non VRAM CUDA dedicata."
    )
    add_metric_table(doc, SYSTEM_ROWS)

    doc.add_heading("3. Stato complessivo e output del run", level=1)
    doc.add_paragraph(
        "run_summary.json viene aggiornato in modo atomico; i CSV vengono appendati "
        "e sincronizzati a ogni stadio, preservando risultati parziali dopo crash."
    )
    add_metric_table(doc, SUMMARY_ROWS)

    chapter = 4
    for title_text, filename, columns in SECTIONS:
        doc.add_heading(f"{chapter}. {title_text}", level=1)
        citation = doc.add_paragraph()
        citation.paragraph_format.space_before = Pt(4)
        citation.paragraph_format.space_after = Pt(4)
        run = citation.add_run(f"Schema stabile: benchmark/{filename}")
        set_run_font(run, size=9, color=MUTED, italic=True)
        rows = [
            (name, unit_for(name), nature_for(name), description_for(name))
            for name in columns
        ]
        add_metric_table(doc, rows)
        if filename == "segmentation_metrics.csv":
            add_callout(
                doc,
                "Interpretazione",
                "coverage_percent, background_percent e unassigned_percent sono "
                "statistiche intrinseche. IoU, Dice, precision, recall e boundary "
                "F-score sono accuratezza solo nelle righe metric_scope=ground_truth.",
            )
        elif filename == "reconstruction_metrics.csv":
            add_callout(
                doc,
                "Protocollo held-out",
                "Le evaluation_indices sono escluse da layer training, global "
                "refinement e baseline monolitica. Scene storiche senza split non "
                "ricevono score post-hoc presentati come held-out.",
            )
        elif filename == "editing_metrics.csv":
            doc.add_heading("Definizioni matematiche dell’editing", level=2)
            doc.add_paragraph(
                "Sia D(x) il cambiamento RGB assoluto medio, T la maschera target, "
                "μT il cambiamento medio interno e μ¬T quello esterno."
            )
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("edit_leakage_ratio = μ¬T / max(μT, 10⁻¹²)")
            set_run_font(r, size=11, color=INK, bold=True)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("edit_locality_score = μT / max(μT + μ¬T, 10⁻¹²)")
            set_run_font(r, size=11, color=INK, bold=True)
            doc.add_paragraph(
                "Leakage più basso e locality più alta sono preferibili. "
                "Disocclusioni reali possono tuttavia produrre variazioni fuori maschera."
            )
        chapter += 1

    doc.add_heading(f"{chapter}. Statistiche aggregate", level=1)
    add_metric_table(doc, [
        ("count", "conteggio", "Derivato", "Numero di osservazioni valide."),
        ("mean", "unità originale", "Derivato", "Media aritmetica."),
        ("median", "unità originale", "Derivato", "Mediana."),
        ("std", "unità originale", "Derivato", "Deviazione standard campionaria."),
        ("min / max", "unità originale", "Derivato", "Estremi osservati."),
        ("p25 / p75", "unità originale", "Derivato", "Primo e terzo quartile."),
        ("ci95_low / ci95_high", "unità originale", "Derivato",
         "Intervallo 95% con approssimazione normale, solo con almeno due valori."),
        ("paired difference", "unità originale", "Derivato",
         "Differenza monolithic − layered per scena."),
        ("paired percent difference", "%", "Derivato",
         "Differenza paired normalizzata sul valore layered."),
        ("success_rate_percent", "%", "Derivato", "Run success / run totali."),
        ("stage_failure_counts", "conteggio", "Derivato", "Fallimenti raggruppati per stadio."),
        ("failure_reasons", "conteggio", "Derivato", "Motivi di fallimento più comuni."),
        ("usable_merged_ply_percent", "%", "Derivato", "Scene con PLY merged valido."),
        ("valid_semantic_layers_percent", "%", "Derivato", "Scene con layer semantici addestrati."),
        ("valid_mood_variant_percent", "%", "Derivato", "Scene con almeno una variante mood."),
    ])
    add_callout(
        doc,
        "Test statistici",
        "Il framework non esegue automaticamente test di significatività, perché "
        "la numerosità campionaria potrebbe non soddisfarne le assunzioni.",
    )

    chapter += 1
    doc.add_heading(f"{chapter}. Metriche attualmente non sempre disponibili", level=1)
    unavailable = [
        "LPIPS e outside-mask LPIPS senza dipendenza opzionale o protocollo mascherato valido.",
        "Thin-structure IoU senza annotazioni manuali dedicate.",
        "Overlap pre-conflitto e motivi storici di discard se le maschere raw non sono persistite.",
        "Metriche held-out per scene già addestrate senza benchmark_view_split.",
        "Editing per istanza quando il PLY non conserva la proprietà intera label.",
        "Target-night ERP PSNR/SSIM senza un riferimento mood renderizzato con protocollo valido.",
        "Accuratezza geometrica novel-view reale senza acquisizioni da camere indipendenti.",
    ]
    for text in unavailable:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        set_run_font(run)

    chapter += 1
    doc.add_heading(f"{chapter}. File prodotti", level=1)
    add_metric_table(doc, [
        ("system_info.json", "JSON", "Metadati", "Sistema, software, Git, comando, seed e ERP."),
        ("experiment_config.json", "JSON", "Configurazione", "Configurazione risolta e scena attiva."),
        ("run_summary.json", "JSON", "Stato", "Esito, durata, failure, output e riepiloghi."),
        ("stage_timings.csv", "CSV", "Raw/derivato", "Tempi strutturati e rate normalizzati."),
        ("resource_samples.csv", "CSV", "Raw", "Serie temporale di memoria, swap e CPU."),
        ("layer_metrics.csv", "CSV", "Raw/derivato", "Una riga per layer."),
        ("segmentation_metrics.csv", "CSV", "Intrinseco/GT", "Copertura e accuratezza opzionale."),
        ("reconstruction_metrics.csv", "CSV", "Derivato", "Una riga per vista held-out."),
        ("rendering_metrics.csv", "CSV", "Derivato", "Latenza/FPS per variante."),
        ("editing_metrics.csv", "CSV", "Derivato", "Località e costo degli edit."),
        ("mood_metrics.csv", "CSV", "Derivato", "Conservazione topologica day/mood."),
        ("failures.json", "JSON", "Diagnostica", "Eccezioni e traceback per scena."),
        ("report/plots/*.png|pdf", "immagine", "Derivato", "Quindici grafici thesis-ready."),
        ("report/report.md", "Markdown", "Report", "Sintesi leggibile e limitazioni scientifiche."),
    ])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "ObjSplat Benchmark Metrics"
    doc.core_properties.subject = "Catalogo delle metriche del framework di benchmark ObjSplat"
    doc.core_properties.author = "ObjSplat"
    doc.core_properties.keywords = "ObjSplat, 3D Gaussian Splatting, benchmark, metriche"
    doc.core_properties.comments = "Generated reproducibly by benchmark/tools/build_metrics_docx.py"
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        default="docs/ObjSplat_Benchmark_Metrics.docx",
        type=Path,
    )
    args = parser.parse_args()
    build_document(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
